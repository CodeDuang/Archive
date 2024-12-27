import sys

from PyQt5.QtWidgets import QApplication, QMainWindow, QFrame, QMessageBox, QFileDialog

import Report.report


class ReportRunMe(QMainWindow, Report.report.Ui_MainWindow):
    def __init__(self):
        super(ReportRunMe, self).__init__()
        self.setupUi(self)
        self.initt()
        self.setWindowTitle("报告")  # 设置标题
        self.setFixedSize(self.width(), self.height())  # 禁止缩放

    def initt(self):
        self.f1.setFrameStyle(1)
        self.f2.setFrameStyle(1)
        self.f3.setFrameStyle(1)
        self.f4.setFrameStyle(1)
        self.f5.setFrameStyle(1)
        self.f6.setFrameStyle(1)
        self.f7.setFrameStyle(1)
        self.f8.setFrameStyle(1)
        self.f9.setFrameStyle(1)
        self.f10.setFrameStyle(1)
        self.f11.setFrameStyle(1)
        self.f12.setFrameStyle(1)
        self.f13.setFrameStyle(1)
        self.f14.setFrameStyle(1)
        # self.gridFrame.setFrameStyle(1)
        # self.horizontalFrame.setFrameStyle(1)
        # self.horizontalFrame_2.setFrameStyle(1)
        self.pushButton.clicked.connect(self.btn_report)  # 生成word报告按钮

        self.data_inject()  # 动态加载数据

    def data_inject(self):
        with open('report.txt', 'r', encoding="utf-8") as f:  # 读取报告信息
            data = f.read().split('\t')

            self.label_2.setText(data[0])  # 报告编号
            if data[1] == '☑仿真芯片数据    ☐实时芯片数据':  # 芯片数据类型
                self.checkBox.setChecked(True)
                self.checkBox_2.setChecked(False)
            else:
                self.checkBox.setChecked(False)
                self.checkBox_2.setChecked(True)
            self.label_5.setText(data[2])
            self.label_7.setText(data[3].split(':')[1])
            if data[4] != '':
                self.label_15.setText(data[4])
            else:
                self.label_15.setText('未检测')
            if data[5] != '':
                self.label_17.setText(data[5])
            else:
                self.label_17.setText('未检测')
            if data[6] != '':
                self.label_19.setText(data[6])
            else:
                self.label_19.setText('未检测')
            if data[7] != '':
                self.label_21.setText(data[7])
            else:
                self.label_21.setText('未检测')

    def btn_report(self):
        # 1.由用户选择对应的目录和文件名
        filepath, type = QFileDialog.getSaveFileName(None, "文件保存", "/",
                                                     'WORD Files(*.docx)')  # 前面是地址，后面是文件类型,得到输入地址的文件名和地址txt(*.txt*.xls);;image(*.png)不同类别
        if filepath == '':
            QMessageBox.warning(None, "Warning", "保存失败，请选择word保存目录和名称")
            return
        # 2.保存文件
        from docx import Document
        from docx.shared import Inches
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.shared import Pt, Cm

        with open('report.txt', 'r', encoding="utf-8") as f:  # 读取报告信息
            data = f.read().split('\t')
            print(data)

            # 新建word对象
            d = Document()

            # 设置word正文文本的格式类型
            d.styles['Normal'].font.name = 'Times New Roman'  # 设置英文文本字体
            d.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文文本字体
            d.styles['Normal'].font.size = Pt(9)  # 设置字体尺寸

            # 添加一个 4*2 表格
            table = d.add_table(rows=8, cols=4, style='Table Grid')
            table.cell(1, 0).width = Inches(1)
            table.cell(1, 1).width = Inches(2)
            table.cell(1, 2).width = Inches(1)
            table.cell(1, 3).width = Inches(1)
            # 表格的第一行
            row = table.rows[0]
            row.height = Cm(2)  # 设置表格行高
            row.cells[0].text = '泄露分析报告'
            row.cells[0].paragraphs[0].runs[0].font.size = Pt(20)  # 设置字体尺寸
            row.cells[0].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER  # 居中文本
            row.cells[0].paragraphs[0].runs[0].font.name = '华文中宋'  # 设置英文文本字体
            row.cells[0].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '华文中宋')  # 设置中文文本字体
            row.cells[0].merge(row.cells[-2])  # 合并单元格，格式：左上角.merge(右下角)
            row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 表格单元格上下居中对齐

            # row.cells[-1].text = '这是报告模板的编号'
            row.cells[-1].text = data[0]  # 报告模板的编号
            row.cells[-1].paragraphs[0].runs[0].font.size = Pt(9)  # 设置字体尺寸
            row.cells[-1].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER  # 居中文本
            row.cells[-1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 表格单元格上下居中对齐

            # 获取单元格（报告类型、芯片ip端口）
            row = table.rows[1]
            row.height = Cm(0.7)  # 设置表格行高
            row = table.rows[2]
            row.height = Cm(0.7)  # 设置表格行高
            cell = table.cell(1, 0)  # 注意，第一个参数表示垂直方向的坐标，第二个参数表示横向坐标
            cell.text = '报告类型：'  # ☐ ☑
            cell.paragraphs[0].runs[0].font.size = Pt(9)  # 设置字体尺寸
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 表格单元格上下居中对齐
            cell = table.cell(1, 1)
            # cell.text = '☐仿真芯片数据    ☑实时芯片数据'
            cell.text = data[1]
            cell.paragraphs[0].runs[0].font.size = Pt(9)  # 设置字体尺寸
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 表格单元格上下居中对齐
            cell = table.cell(2, 0)
            cell.text = '芯片ip端口：'
            cell.paragraphs[0].runs[0].font.size = Pt(9)  # 设置字体尺寸
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 表格单元格上下居中对齐
            cell = table.cell(2, 1)
            # cell.text = '127.0.0.1:8000'
            cell.text = data[2]  # ip
            cell.paragraphs[0].runs[0].font.size = Pt(9)  # 设置字体尺寸
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 表格单元格上下居中对齐

            # 合并单元格(报告完成时间)
            table.cell(1, 2).merge(table.cell(2, 3))
            # table.cell(1, 2).text = '报告完成时间：2000年1月1日'
            table.cell(1, 2).text = data[3]  # 报告完成时间
            table.cell(1, 2).paragraphs[0].runs[0].font.size = Pt(9)  # 设置字体尺寸
            table.cell(1, 2).paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER  # 居中文本
            table.cell(1, 2).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 表格单元格上下居中对齐

            # 侧信道分析方法和结论部分
            row = table.rows[3]
            row.height = Cm(1)  # 设置表格行高
            way = ['侧信道分析方法', 'DPA(差分功耗分析)', 'TA(模板分析)', 'CPA(相关性功耗分析)', 'SPA(简单功耗分析)']
            conclusion = ['结论', '未检测', '未检测', '未检测', '未检测', '未检测', '未检测']
            for i in range(3, 8):  # 遍历行(侧行道分析方法和结论)
                row = table.rows[i]
                row.cells[1].merge(row.cells[-1])  # 合并单元格
                row.cells[0].text = way[i - 3]  # 侧信道分析方法
                row.cells[0].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER  # 居中文本
                row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 表格单元格上下居中对齐
                if i > 3 and data[i] != '':
                    row.cells[1].text = data[i]

                else:
                    row.cells[1].text = conclusion[i - 3]  # 侧信道分析结果
                row.cells[1].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER  # 居中文本
                row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 表格单元格上下居中对齐
                if i != 3:
                    row.height = Cm(4)
            row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 表格单元格上下居中对齐

            # 保存
            d.save(filepath)  # 注意：字符串前面加r是防止字符转义
            QMessageBox.information(None, "Information", "word报告已生成！")


if __name__ == "__main__":
    app = QApplication(sys.argv)

    # 定义页面对象
    reportt = ReportRunMe()  # 综合导航栏界面

    reportt.show()
    sys.exit(app.exec_())
