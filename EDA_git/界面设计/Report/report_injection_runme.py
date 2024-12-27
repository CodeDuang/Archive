import sys

from PyQt5.QtWidgets import QApplication, QMainWindow, QMessageBox, QTableWidgetItem, QFileDialog, QAbstractItemView

import Report.report_injection


class ReportInjectionRunMe(QMainWindow, Report.report_injection.Ui_MainWindow):
    def __init__(self):
        super(ReportInjectionRunMe, self).__init__()
        self.setupUi(self)
        self.initt()
        self.setWindowTitle("报告")  # 设置标题
        self.setFixedSize(self.width(), self.height())  # 禁止缩放

    def initt(self):
        # 给边框加上线条
        self.f1.setFrameStyle(1)
        self.f2.setFrameStyle(1)
        self.f3.setFrameStyle(1)
        self.f4.setFrameStyle(1)
        self.f5.setFrameStyle(1)
        self.f6.setFrameStyle(1)

        self.textBrowser.setStyleSheet('border-width:0;border-style:outset')  # textBrowser设置无边框

        self.tableWidget_2.setSelectionMode(QAbstractItemView.NoSelection) #设置点击模式为不可点击

        self.pushButton.clicked.connect(self.btn_word)  # 生成word按钮

        self.data_inject()  # 动态加载数据

    def data_inject(self):
        with open('report_injection.txt', 'r', encoding="utf-8") as f:  # 读取报告信息
            data = f.read().split('\n')
            self.label_2.setText(data[0])  # 报告编号
            self.label_5.setText(data[2])  # ip端口
            self.label_7.setText(data[3].split(':')[1])  # 完成报告时间
            self.textBrowser.setText('\n'.join(data[5:]))  # 故障结果

            data_inject = data[4].split(' ')  # 故障信息变成数组
            lon = len(data_inject) // 10  # 获取注入信息行数
            i = 0  # 用来遍历data_inject数组的
            for row in range(lon):  # 按行写入故障信息
                # 判断是否需要加一行（扩充逻辑）
                if self.tableWidget_2.rowCount() <= row:
                    self.tableWidget_2.setRowCount(row + 1)  # 设置行数
                # 故障注入部分，写入items数据
                for column in range(self.tableWidget_2.columnCount()):
                    self.tableWidget_2.setItem(row, column, QTableWidgetItem(data_inject[i]))
                    i += 1
        f.close()

    # 在word中生成word格式的报告
    def btn_word(self):
        # 1.由用户选择对应的目录和文件名
        filepath, type = QFileDialog.getSaveFileName(None, "文件保存", "/",
                                                     'WORD Files(*.docx)')  # 前面是地址，后面是文件类型,得到输入地址的文件名和地址txt(*.txt*.xls);;image(*.png)不同类别
        if filepath == '':
            QMessageBox.warning(None, "Warning", "保存失败，请选择word保存目录和名称")
            return

        # 2.开始保存
        from docx import Document
        from docx.shared import Inches
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.shared import Pt, Cm

        with open('report_injection.txt', 'r', encoding="utf-8") as f:  # 读取报告信息
            data = f.read().split('\n')
            print(data)

            # 新建word对象
            d = Document()

            # 设置word正文文本的格式类型
            d.styles['Normal'].font.name = 'Times New Roman'  # 设置英文文本字体
            d.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文文本字体
            d.styles['Normal'].font.size = Pt(14)  # 设置字体尺寸

            # 添加一个 4*2 表格
            table = d.add_table(rows=5, cols=4, style='Table Grid')
            table.cell(1, 0).width = Inches(1)
            table.cell(1, 1).width = Inches(2)
            table.cell(1, 2).width = Inches(1)
            table.cell(1, 3).width = Inches(1)
            # 表格的第一行
            row = table.rows[0]
            row.height = Cm(2)  # 设置表格行高
            row.cells[0].text = '故障注入报告'
            row.cells[0].paragraphs[0].runs[0].font.size = Pt(20)  # 设置字体尺寸
            row.cells[0].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER  # 居中文本
            row.cells[0].paragraphs[0].runs[0].font.name = '华文中宋'  # 设置英文文本字体
            row.cells[0].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '华文中宋')  # 设置中文文本字体
            row.cells[0].merge(row.cells[-2])  # 合并单元格，格式：左上角.merge(右下角)
            row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 表格单元格上下居中对齐

            # row.cells[-1].text = '这是报告模板的编号'
            row.cells[-1].text = data[0]  # 报告模板的编号
            row.cells[-1].paragraphs[0].runs[0].font.size = Pt(9)  # 设置字体尺寸
            row.cells[-1].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER  # 居中文本
            row.cells[-1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 表格单元格上下居中对齐

            # 获取单元格（报告类型、芯片ip端口）
            cell = table.cell(1, 0)  # 注意，第一个参数表示垂直方向的坐标，第二个参数表示横向坐标
            cell.text = '报告类型：'  # ☐ ☑
            cell.paragraphs[0].runs[0].font.size = Pt(9)  # 设置字体尺寸
            cell = table.cell(1, 1)
            # cell.text = '☐仿真芯片数据    ☑实时芯片数据'
            cell.text = data[1]
            cell.paragraphs[0].runs[0].font.size = Pt(9)  # 设置字体尺寸
            cell = table.cell(2, 0)
            cell.text = '芯片ip端口：'
            cell.paragraphs[0].runs[0].font.size = Pt(9)  # 设置字体尺寸
            cell = table.cell(2, 1)
            # cell.text = '127.0.0.1:8000'
            cell.text = data[2]  # ip
            cell.paragraphs[0].runs[0].font.size = Pt(9)  # 设置字体尺寸

            # 合并单元格(报告完成时间)
            table.cell(1, 2).merge(table.cell(2, 3))
            # table.cell(1, 2).text = '报告完成时间：2000年1月1日'
            table.cell(1, 2).text = data[3]  # 报告完成时间
            table.cell(1, 2).paragraphs[0].runs[0].font.size = Pt(9)  # 设置字体尺寸
            table.cell(1, 2).paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER  # 居中文本
            table.cell(1, 2).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 表格单元格上下居中对齐

            # 合并故障分析信息单元格
            row = table.rows[3]
            row.cells[0].merge(row.cells[-1])
            row.height = Cm(6)  # 设置表格行高
            row.cells[0].text = '故障注入信息'
            row.cells[0].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER  # 居中文本

            # 开始读取故障信息，并且根据故障信息生成表格，并装载故障信息
            dataa = data[4].split(' ')  # 先把故障信息放入dataa列表中
            dataa_len = len(dataa)  # 获取列表长度，以计算需要生成几行表格
            ttable = row.cells[0].add_table(rows=(dataa_len // 10) + 1, cols=10)  # 单元格中加表格（按照故障信息长度生成表格）
            ttable.style = 'Table Grid'
            table_title = ['所属模块', '寄存器名称', '起始ID', '控制时钟', '终止ID', '注入起始时刻', '故障类型', '故障注入间隔', '故障次数', '故障持续周期']
            for i in range(10):  # 故障信息第一行为列标题
                ttable.cell(0, i).text = table_title[i]
            for i in range(dataa_len // 10):  # 故障信息装载
                for j in range(10):
                    ttable.cell(i + 1, j).text = dataa[i * 10 + j]

            for row in ttable.rows:  # 更改表格内全部字体大小
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(8)

            # 故障分析结果部分
            row = table.rows[4]
            row.cells[0].text = '故障分析结果'
            row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 表格单元格上下居中对齐

            row.cells[1].merge(row.cells[-1])  # 合并单元格，格式：左上角.merge(右下角)
            row.height = Cm(5)  # 设置表格行高
            for txt in data[5:]:  # 故障结果注入
                row.cells[1].add_paragraph(txt)
            for paragraph in row.cells[1].paragraphs:  # 遍历更改故障结果的字体大小
                for run in paragraph.runs:
                    run.font.size = Pt(8)
            row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 表格单元格上下居中对齐

            # 保存
            d.save(filepath)  # 注意：字符串前面加r是防止字符转义
            QMessageBox.information(None, "Information", "word报告已生成！")
        f.close()


if __name__ == "__main__":
    app = QApplication(sys.argv)

    # 定义页面对象
    reportt = ReportInjectionRunMe()  # 综合导航栏界面

    reportt.show()
    sys.exit(app.exec_())
